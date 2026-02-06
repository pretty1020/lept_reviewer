"""
LEPT AI Reviewer - Document Processing Service
"""

import io
from typing import Optional, Tuple
from pathlib import Path

import streamlit as st


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[bool, str]:
    """
    Extract text from a PDF file.
    
    Args:
        file_bytes: PDF file as bytes
    
    Returns:
        Tuple of (success, extracted_text_or_error)
    """
    try:
        from PyPDF2 import PdfReader
        
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        if not text_parts:
            return False, "Could not extract text from the PDF. The file may contain only images."
        
        full_text = "\n\n".join(text_parts)
        
        # Clean up the text
        full_text = clean_extracted_text(full_text)
        
        if len(full_text.strip()) < 50:
            return False, "Extracted text is too short. Please upload a document with more content."
        
        return True, full_text
        
    except Exception as e:
        return False, f"Error reading PDF: {str(e)}"


def extract_text_from_docx(file_bytes: bytes) -> Tuple[bool, str]:
    """
    Extract text from a DOCX file.
    
    Args:
        file_bytes: DOCX file as bytes
    
    Returns:
        Tuple of (success, extracted_text_or_error)
    """
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        if not text_parts:
            return False, "Could not extract text from the DOCX file."
        
        full_text = "\n\n".join(text_parts)
        
        # Clean up the text
        full_text = clean_extracted_text(full_text)
        
        if len(full_text.strip()) < 50:
            return False, "Extracted text is too short. Please upload a document with more content."
        
        return True, full_text
        
    except Exception as e:
        return False, f"Error reading DOCX: {str(e)}"


def extract_text_from_file(uploaded_file) -> Tuple[bool, str]:
    """
    Extract text from an uploaded file (PDF or DOCX).
    
    Args:
        uploaded_file: Streamlit UploadedFile object
    
    Returns:
        Tuple of (success, extracted_text_or_error)
    """
    if uploaded_file is None:
        return False, "No file provided"
    
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()
    
    # Reset file pointer for potential reuse
    uploaded_file.seek(0)
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    else:
        return False, "Unsupported file format. Please upload a PDF or DOCX file."


def clean_extracted_text(text: str) -> str:
    """
    Clean up extracted text.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Strip whitespace
        line = line.strip()
        
        # Skip empty lines (but keep some for paragraph separation)
        if not line:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
        
        # Remove lines that are just special characters
        if all(c in '.-_=*#' for c in line.replace(' ', '')):
            continue
        
        cleaned_lines.append(line)
    
    cleaned_text = '\n'.join(cleaned_lines)
    
    # Remove multiple consecutive newlines (keep max 2)
    while '\n\n\n' in cleaned_text:
        cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
    
    return cleaned_text.strip()


def truncate_text_for_ai(text: str, max_chars: int = 15000) -> str:
    """
    Truncate text to fit within AI context limits.
    
    Args:
        text: Full text
        max_chars: Maximum character count
    
    Returns:
        Truncated text
    """
    if len(text) <= max_chars:
        return text
    
    # Try to truncate at a sentence boundary
    truncated = text[:max_chars]
    last_period = truncated.rfind('.')
    
    if last_period > max_chars * 0.8:
        return truncated[:last_period + 1]
    
    return truncated + "..."


def get_text_stats(text: str) -> dict:
    """
    Get statistics about extracted text.
    
    Args:
        text: Extracted text
    
    Returns:
        Dictionary with text statistics
    """
    if not text:
        return {
            "char_count": 0,
            "word_count": 0,
            "line_count": 0,
            "paragraph_count": 0
        }
    
    words = text.split()
    lines = text.split('\n')
    paragraphs = [p for p in text.split('\n\n') if p.strip()]
    
    return {
        "char_count": len(text),
        "word_count": len(words),
        "line_count": len(lines),
        "paragraph_count": len(paragraphs)
    }
